from docx import Document
from docx.shared import Inches
import numpy as np
import json
import matplotlib.pyplot as plt
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

class ReportCreator():
    def __init__(self,paras) -> None:
        self.title = "AEL Results"
        self.exp_output_path = paras.exp_output_path
        self.ec_n_pop = paras.ec_n_pop
        self.ec_pop_size = paras.ec_pop_size
        self.paras = paras

    def create_convergence(self):
        n_start = 0
        obj_list = np.zeros((self.ec_n_pop,self.ec_pop_size))
        for i in range(n_start,n_start+self.ec_n_pop):
            ### Get result ###
            #Load JSON data from file
            with open(self.exp_output_path+"ael_results/pops/population_generation_"+str(i)+".json") as file:
                    data = json.load(file)


            #Print each individual in the population
            na = 0
            for individual in data:
                code = individual['code']
                alg = individual['algorithm']
                obj = individual['objective']
                
                #code2file(alg,code,na,i)
                #print(obj)
                obj_list[i-n_start,na] = obj
                na +=1

        # Set font family to Times New Roman
        plt.rcParams['font.family'] = 'Times New Roman'
        # Generate x-axis values for number of generations
        generations = np.arange(1, obj_list.shape[0] + 1)
        best_objective = np.min(obj_list, axis=1)
        mean_objective = np.mean(obj_list, axis=1)

        # Set figure size
        plt.figure(figsize=(10, 6), dpi=80)

        # Plot objective value vs. number of generations for all samples as scatter points
        for i in generations:
            plt.scatter(i*np.ones(self.ec_pop_size), obj_list[i-1, :], color='tab:blue', alpha=0.6,s=200)

        # Plot mean and best objectives
        plt.plot(generations, mean_objective, label='Mean', color='orange',linewidth=3.0)
        plt.plot(generations, best_objective, label='Best', color='r',linewidth=3.0)

        # Set plot title and labels with enlarged font size
        plt.xlabel('Number of Generations', fontsize=18)
        plt.ylabel('Obj.', fontsize=20)

        objmin = np.min(obj_list)
        objmax = np.max(obj_list)
        delta=(objmax-objmin)/100.0
        # Set y-axis range
        plt.ylim([objmin-delta, objmax+delta])

        # Add scatter legend with enlarged font size
        plt.scatter([], [], color='tab:blue', alpha=0.6, label='Algorithms',s=200)  # Empty scatter plot for legend
        #plt.legend(scatterpoints=1, frameon=False, labelspacing=1, fontsize=20)
        plt.legend(scatterpoints=1, frameon=True, labelspacing=1, fontsize=20, fancybox=True, facecolor='gainsboro')
        # Adjust ticks and grid
        plt.xticks(np.arange(1, obj_list.shape[0] + 1, 2),fontsize=18)
        plt.yticks(np.arange(objmin-delta, objmax+delta, 10),fontsize=18)
        plt.grid(True, which='both', linestyle='--', linewidth=0.5)

        # Show the plot
        plt.tight_layout()
        plt.savefig(self.exp_output_path+'ael_results/ael_convergence.png')   # Save the plot as a file
        plt.savefig(self.exp_output_path+'ael_results/ael_convergence.pdf') 
        #plt.show()
    
    def get_final_algorithms(self):
        ### Get result ###
        with open(self.exp_output_path+"ael_results/pops/population_generation_"+str(self.ec_n_pop)+".json") as file:
            data = json.load(file)

        # for individual in data:
        #     #print(individual)
        #     results = individual
        #     code = results['code']
        #     algorithm = results['algorithm']
        #     gap = results['objective']
            
        #     #code2file(code)
            
        #     print("### algorithm: \n",algorithm)
        #     print("### code: \n",code)
        #     print("### Average gap is : \n",gap)
        #     input() 
        return data      
    


    def generate_doc_report(self):
        # Create a new Document
        doc = Document()

        # Add Title
        doc.add_heading(self.title, level=1)

        # Add Parameter Settings
        doc.add_heading('Parameter Settings', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set Table Style
        table.style = 'Table Grid'

        # Set Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True  # Make headers bold

        # Add Data to Table
        for attr, value in vars(self.paras).items():
            row_cells = table.add_row().cells
            row_cells[0].text = attr
            row_cells[1].text = str(value)

        # Add Borderlines
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)  # Set font size
                        run.font.bold = False  # Set font style
                        run.font.name = 'Calibri'  # Set font type
                        run.font.color.rgb = RGBColor(0, 0, 0) # Set font color



        # Add Convergence Process
        doc.add_heading('Convergence Process', level=2)
        self.create_convergence()
        doc.add_picture(self.exp_output_path+'ael_results/ael_convergence.png', width=Inches(4))

        # Add Final Results
        doc.add_heading('Final Results', level=2)
        algorithms_data = self.get_final_algorithms()

        # Add top five algorithms data
        doc.add_heading('Top Five Algorithms', level=3)
        for i, algorithm_data in enumerate(algorithms_data[:5]):
            doc.add_heading(f'Algorithm {i+1}', level=4)
            doc.add_paragraph(f'Algorithm: {algorithm_data["algorithm"]}')
                    # Create a new paragraph
            p = doc.add_paragraph()

            # Add the code block with background color and border
            code = algorithm_data["code"]
            code_block = p.add_run()
            code_block.text = f'Code:\n{code}'
            code_block_font = code_block.font
            code_block_font.size = Pt(8)

            # Set the background color
            shading_elm = OxmlElement('w:shd')
            shading_elm.set('fill', 'D9D9D9')  # Set your desired background color here
            p._element.append(shading_elm)

            # Set the border
            #p.border_top.space = Pt(1)     # Set border space

            doc.add_paragraph(f'Fitness: {algorithm_data["objective"]}')
            doc.add_paragraph('')  # Add a blank paragraph for separation

        # Save the document
        doc.save('ael_report.docx')



if __name__ == "__main__":

    Paras = Paras()

    RC = ReportCreator(Paras)

    RC.generate_doc_report()

    print("Doc report generated successfully!")